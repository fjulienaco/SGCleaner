import streamlit as st
import os
import tempfile
import re
from typing import List, Dict, Tuple
import io
import json

# Handle optional imports gracefully
try:
    import markdown
except ImportError:
    markdown = None

try:
    import openai
except ImportError:
    openai = None

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    def load_dotenv():
        pass

# Handle docx import with better error handling
try:
    from docx import Document
except ImportError as e:
    st.error(f"""
    **Missing Required Package**: python-docx is not installed.
    
    Error: {e}
    
    Please install it with: `pip install python-docx`
    
    If you're seeing this on Streamlit Community Cloud, please check the requirements.txt file.
    """)
    st.stop()

# Load environment variables
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="Style Guide Cleaner Pro",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

def setup_openai():
    """Setup OpenAI client with API key"""
    if openai is None:
        return None
        
    # Try multiple sources for API key
    api_key = (
        st.session_state.get('openai_api_key') or  # User input in sidebar
        os.getenv('OPENAI_API_KEY') or  # Environment variable
        st.secrets.get('OPENAI_API_KEY')  # Streamlit Community Cloud secrets
    )
    
    # Debug information (remove quotes if present)
    if api_key and api_key.startswith("'") and api_key.endswith("'"):
        api_key = api_key[1:-1]
    elif api_key and api_key.startswith('"') and api_key.endswith('"'):
        api_key = api_key[1:-1]
    
    if not api_key:
        return None
    
    try:
        client = openai.OpenAI(api_key=api_key)
        return client
    except Exception as e:
        st.error(f"Error setting up OpenAI: {e}")
        return None

def extract_text_from_docx(doc) -> Dict[str, any]:
    """Extract text content from docx document with enhanced analysis"""
    content = {
        'paragraphs': [],
        'tables': [],
        'headers': [],
        'footers': [],
        'metadata': {
            'total_paragraphs': 0,
            'total_tables': 0,
            'styles_used': set(),
            'complexity_indicators': []
        }
    }
    
    # Extract paragraphs with enhanced analysis
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text = paragraph.text.strip()
            style = paragraph.style.name if paragraph.style else 'Normal'
            
            para_info = {
                'text': text,
                'style': style,
                'length': len(text),
                'complexity_score': calculate_complexity_score(text),
                'cleaning_opportunities': identify_cleaning_opportunities(text)
            }
            
            content['paragraphs'].append(para_info)
            content['metadata']['styles_used'].add(style)
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            if row_data:
                table_data.append(row_data)
        if table_data:
            content['tables'].append(table_data)
    
    content['metadata']['total_paragraphs'] = len(content['paragraphs'])
    content['metadata']['total_tables'] = len(content['tables'])
    
    return content

def calculate_complexity_score(text: str) -> float:
    """Calculate complexity score for a text (0-1, higher = more complex)"""
    score = 0.0
    
    # Length factor
    if len(text) > 200:
        score += 0.2
    if len(text) > 500:
        score += 0.2
    
    # Formatting complexity
    if re.search(r'[.]{3,}|[-_]{3,}', text):
        score += 0.1
    
    # Reference complexity
    if re.search(r'\b(?:see|refer to|page|section)\s+', text, re.IGNORECASE):
        score += 0.1
    
    # List/numbering complexity
    if re.search(r'^\s*\d+[\.\)]\s*', text, re.MULTILINE):
        score += 0.1
    
    # Version/date references
    if re.search(r'\b(?:version|v\.?)\s*\d+|date|updated\b', text, re.IGNORECASE):
        score += 0.1
    
    # Human workflow elements (high priority for cleaning)
    if re.search(r'\b(?:consult|check|discuss|send|get approval|escalate|verify|double-check)\s+(?:with|to)\s+(?:project manager|PM|team lead|supervisor|linguist|client|reviewer)', text, re.IGNORECASE):
        score += 0.3
    
    if re.search(r'\b(?:email|call|contact|notify)\s+(?:the team|client|manager|supervisor)', text, re.IGNORECASE):
        score += 0.2
    
    if re.search(r'\b(?:log in|update|enter|submit)\s+(?:the system|database|platform)', text, re.IGNORECASE):
        score += 0.2
    
    if re.search(r'\b(?:quality assurance|QA|review|approval)\s+(?:process|workflow|step)', text, re.IGNORECASE):
        score += 0.2
    
    return min(score, 1.0)

def identify_cleaning_opportunities(text: str) -> List[str]:
    """Identify specific cleaning opportunities in text"""
    opportunities = []
    
    if re.search(r'\b(?:page|p\.?)\s*\d+\b', text, re.IGNORECASE):
        opportunities.append('page_reference')
    
    if re.search(r'\b(?:version|v\.?)\s*\d+\.?\d*\b', text, re.IGNORECASE):
        opportunities.append('version_reference')
    
    if re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text):
        opportunities.append('date_reference')
    
    if re.search(r'\b(?:document|doc|file)\s*(?:name|title|id)\s*:.*?(?=\n|\r|$)', text, re.IGNORECASE):
        opportunities.append('metadata_reference')
    
    if re.search(r'[.]{3,}|[-_]{3,}', text):
        opportunities.append('excessive_punctuation')
    
    if len(text) > 300 and re.search(r'\b(?:also|furthermore|moreover|additionally)\b', text, re.IGNORECASE):
        opportunities.append('wordy_text')
    
    # Human workflow elements (not useful for LLMs)
    if re.search(r'\b(?:consult|check|discuss|send|get approval|escalate|verify|double-check)\s+(?:with|to)\s+(?:project manager|PM|team lead|supervisor|linguist|client|reviewer)', text, re.IGNORECASE):
        opportunities.append('human_workflow')
    
    if re.search(r'\b(?:email|call|contact|notify)\s+(?:the team|client|manager|supervisor)', text, re.IGNORECASE):
        opportunities.append('communication_protocols')
    
    if re.search(r'\b(?:log in|update|enter|submit)\s+(?:the system|database|platform)', text, re.IGNORECASE):
        opportunities.append('administrative_procedures')
    
    if re.search(r'\b(?:quality assurance|QA|review|approval)\s+(?:process|workflow|step)', text, re.IGNORECASE):
        opportunities.append('human_qa_workflow')
    
    if re.search(r'\b(?:before|after|during)\s+(?:translation|project)\s+(?:consult|check|discuss|send)', text, re.IGNORECASE):
        opportunities.append('process_instructions')
    
    # Additional human workflow patterns
    if re.search(r'\b(?:don\'t hesitate to|please don\'t hesitate to)\s+(?:get in touch with|contact)', text, re.IGNORECASE):
        opportunities.append('contact_instructions')
    
    if re.search(r'\b(?:please notify|notify your project manager)', text, re.IGNORECASE):
        opportunities.append('notification_instructions')
    
    if re.search(r'\b(?:queries are welcome|if anything is unclear)', text, re.IGNORECASE):
        opportunities.append('query_instructions')
    
    if re.search(r'https?://[^\s]+|www\.[^\s]+', text):
        opportunities.append('urls_links')
    
    return opportunities

def clean_text_for_llm(text: str, aggressive: bool = False, options: Dict = None) -> str:
    """Enhanced text cleaning with configurable options"""
    if options is None:
        options = {}
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove page numbers and references (if enabled)
    if options.get('remove_page_refs', True):
        text = re.sub(r'\b(?:page|p\.?)\s*\d+\b', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:see|refer to)\s+(?:page|section|chapter)\s*\d+\b', '', text, flags=re.IGNORECASE)
    
    # Remove version numbers and dates (if enabled)
    if options.get('remove_version_refs', True):
        text = re.sub(r'\b(?:version|v\.?)\s*\d+\.?\d*\b', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', '', text)
    
    # Remove document metadata references (if enabled)
    if options.get('remove_metadata', True):
        text = re.sub(r'\b(?:document|doc|file)\s*(?:name|title|id)\s*:.*?(?=\n|\r|$)', '', text, flags=re.IGNORECASE)
    
    # Remove URLs and web links (if enabled)
    if options.get('remove_urls', True):
        text = re.sub(r'https?://[^\s]+', '', text)
        text = re.sub(r'www\.[^\s]+', '', text)
    
    # Clean bullet points and numbering
    text = re.sub(r'^[\s]*[•·▪▫‣⁃]\s*', '- ', text, flags=re.MULTILINE)
    text = re.sub(r'^[\s]*\d+[\.\)]\s*', '', text, flags=re.MULTILINE)
    
    # Remove excessive punctuation (if enabled)
    if options.get('remove_excessive_punctuation', True):
        text = re.sub(r'[.]{3,}', '...', text)
        text = re.sub(r'[-_]{3,}', '---', text)
    
    # Remove wordy text (if enabled)
    if options.get('remove_wordy_text', True):
        text = re.sub(r'\b(?:please note|it is important to note|it should be noted)\b[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:this document|this guide|this style guide)\b[^.]*\.', '', text, flags=re.IGNORECASE)
    
    # Remove human workflow elements (if enabled)
    if options.get('remove_contact_instructions', True):
        text = re.sub(r'\b(?:consult|check|discuss|send|get approval|escalate|verify|double-check)\s+(?:with|to)\s+(?:project manager|PM|team lead|supervisor|linguist|client|reviewer)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:email|call|contact|notify)\s+(?:the team|client|manager|supervisor)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:don\'t hesitate to|please don\'t hesitate to)\s+(?:get in touch with|contact)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:please notify|notify your project manager)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:queries are welcome|if anything is unclear)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:they will provide|forward the query)[^.]*\.', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(?:document the issues|ask you to do so)[^.]*\.', '', text, flags=re.IGNORECASE)
    
    # Remove QA workflows (if enabled)
    if options.get('remove_qa_workflows', True):
        text = re.sub(r'\b(?:quality assurance|QA|review|approval)\s+(?:process|workflow|step)[^.]*\.', '', text, flags=re.IGNORECASE)
    
    # Remove admin procedures (if enabled)
    if options.get('remove_admin_procedures', True):
        text = re.sub(r'\b(?:log in|update|enter|submit)\s+(?:the system|database|platform)[^.]*\.', '', text, flags=re.IGNORECASE)
    
    # Remove human workflow sections (if enabled)
    if options.get('remove_human_workflow', True):
        if re.search(r'\b(?:queries|inconsistencies|feedback|useful links)\b', text, re.IGNORECASE):
            if len(text) < 200:  # Short sections that are likely just about human workflow
                text = ""
    
    if aggressive:
        # More aggressive cleaning - apply additional rules
        text = re.sub(r'\b(?:before|after|during)\s+(?:translation|project)\s+(?:consult|check|discuss|send)[^.]*\.', '', text, flags=re.IGNORECASE)
    
    return text.strip()

def is_human_workflow_section(text: str) -> bool:
    """Check if a section is purely about human workflow and should be removed"""
    text_lower = text.lower()
    
    # Check for section headers that are clearly human workflow
    workflow_headers = [
        'queries', 'inconsistencies', 'feedback', 'useful links',
        'project manager', 'contact', 'notify', 'escalate',
        'change history', 'table of contents'
    ]
    
    for header in workflow_headers:
        if header in text_lower and len(text) < 300:  # Short sections with these headers
            return True
    
    # Check for specific problematic phrases that indicate human workflow
    workflow_phrases = [
        'notify your project manager',
        'contact your project manager', 
        'don\'t hesitate to get in touch',
        'queries are welcome',
        'forward the query',
        'document the issues',
        'ask you to do so',
        'they will provide',
        'escalate to supervisor',
        'send to reviewer',
        'get approval from client',
        'please consult',
        'please also note'
    ]
    
    for phrase in workflow_phrases:
        if phrase in text_lower:
            return True
    
    # Check if the text is mostly about contacting people
    contact_words = ['contact', 'notify', 'email', 'call', 'discuss', 'escalate', 'forward', 'provide', 'consult']
    contact_count = sum(1 for word in contact_words if word in text_lower)
    
    if contact_count >= 2 and len(text) < 200:  # Multiple contact references in short text
        return True
    
    # Check for table of contents entries (page numbers)
    if re.search(r'\d+\s*$', text.strip()):  # Ends with a number (page reference)
        return True
    
    # Check for very short text that's likely metadata
    if len(text.strip()) < 50 and any(word in text_lower for word in ['porsche', 'translation', 'style', 'guide']):
        return True
    
    return False

def optimize_with_openai(client, text: str, context: str = "") -> str:
    """Use OpenAI to optimize text for LLM consumption, removing human workflow elements"""
    try:
        # Check if aggressive mode is enabled
        is_aggressive = "AGGRESSIVE MODE" in context
        
        prompt = f"""
You are an expert at optimizing style guide content for AI language models. Your task is to {"EXTREMELY AGGRESSIVELY" if is_aggressive else "aggressively"} clean this text to be perfect for LLM translation by removing ALL human workflow elements.

Context: {context}

Original text:
{text}

{"EXTREMELY AGGRESSIVELY" if is_aggressive else "AGGRESSIVELY"} REMOVE (these are NOT useful for LLMs):
- Process instructions for human linguists (e.g., "consult with project manager", "check with team lead")
- Human workflow steps (e.g., "send to reviewer", "get approval from client", "notify your project manager")
- Project management references (e.g., "discuss with PM", "escalate to supervisor", "contact your project manager")
- Human communication protocols (e.g., "email the team", "call the client", "get in touch with")
- Manual verification steps (e.g., "double-check with linguist", "verify with human")
- Administrative procedures (e.g., "log in the system", "update the database")
- Quality assurance workflows involving humans
- Any references to human roles, titles, or organizational structure for process purposes
- Queries sections that are about contacting humans
- Feedback and inconsistency reporting instructions
- Links to external websites (unless they contain essential translation resources)
- Entire sections that are only about human workflow processes

KEEP and OPTIMIZE:
- Translation rules and guidelines
- Terminology preferences
- Style conventions
- Grammar rules
- Cultural considerations
- Brand voice guidelines
- Technical specifications for translation
- Content-specific instructions
- Language-specific formatting rules

SPECIAL INSTRUCTIONS:
- If the entire text is about human workflow (like "Queries", "Inconsistencies and feedback", "Useful links"), return "REMOVE_SECTION"
- If the text contains mostly human workflow elements, be very aggressive in removing them
- Remove URLs and web links unless they are essential translation resources
- Remove instructions about contacting people or reporting issues

Make the result:
1. Direct and actionable for AI translation ONLY
2. Focused purely on translation quality and style
3. Completely free of human workflow elements
4. Concise and clear for LLM understanding
5. If nothing useful remains after cleaning, return "REMOVE_SECTION"

Rewritten text:"""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.1
        )
        
        result = response.choices[0].message.content.strip()
        
        # Check if the AI wants to remove the entire section
        if "REMOVE_SECTION" in result:
            return ""
        
        return result
    
    except Exception as e:
        st.warning(f"OpenAI optimization failed: {e}")
        return text

def convert_to_markdown(content: Dict, use_openai: bool = False, client=None, options: Dict = None, ai_aggressive: bool = False) -> str:
    """Convert content to markdown with optional OpenAI optimization"""
    markdown_content = []
    
    # Add title
    markdown_content.append("# Style Guide\n")
    
    # Process paragraphs with section tracking
    i = 0
    while i < len(content['paragraphs']):
        para = content['paragraphs'][i]
        text = clean_text_for_llm(para['text'], aggressive=False, options=options)
        
        # Check if this is a heading
        style = para['style'].lower()
        is_heading = 'heading' in style or 'title' in style
        
        if is_heading:
            # This is a heading - check if the following content should be kept
            section_should_be_removed = False
            
            # Look ahead to see if the section content should be removed
            j = i + 1
            while j < len(content['paragraphs']):
                next_para = content['paragraphs'][j]
                next_style = next_para['style'].lower()
                next_is_heading = 'heading' in next_style or 'title' in next_style
                
                # If we hit another heading, stop looking
                if next_is_heading:
                    break
                
                # Check if the content after this heading should be removed
                next_text = clean_text_for_llm(next_para['text'], aggressive=False, options=options)
                if (options.get('remove_human_workflow', True) and is_human_workflow_section(next_text)) or not next_text:
                    section_should_be_removed = True
                    break
                
                j += 1
            
            # If the section should be removed, skip the heading and its content
            if section_should_be_removed:
                # Skip this heading and all content until next heading
                while i < len(content['paragraphs']):
                    current_para = content['paragraphs'][i]
                    current_style = current_para['style'].lower()
                    if 'heading' in current_style or 'title' in current_style:
                        if i > 0:  # Don't skip the first heading we're already processing
                            i += 1
                            break
                    i += 1
                continue
        
        # Process regular content (non-heading or heading that should be kept)
        if not text:
            i += 1
            continue
        
        # Pre-filter: Skip entire sections that are clearly human workflow (if enabled)
        if options.get('remove_human_workflow', True) and is_human_workflow_section(text):
            i += 1
            continue
        
        # Apply OpenAI optimization if enabled and client available
        if use_openai and client and (para['complexity_score'] > 0.2 or len(para['cleaning_opportunities']) > 0):
            context = f"Style: {para['style']}, Complexity: {para['complexity_score']:.2f}"
            if ai_aggressive:
                context += ", AGGRESSIVE MODE"
            text = optimize_with_openai(client, text, context)
            
            # If AI optimization removed the content, skip this paragraph
            if not text or text.strip() == "":
                i += 1
                continue
        
        # Handle different paragraph styles
        if is_heading:
            level = 1
            if '1' in style:
                level = 1
            elif '2' in style:
                level = 2
            elif '3' in style:
                level = 3
            else:
                level = 2
            
            markdown_content.append(f"{'#' * level} {text}\n")
        else:
            # Regular paragraph
            markdown_content.append(f"{text}\n")
        
        i += 1
    
    # Process tables
    if content['tables']:
        markdown_content.append("\n## Tables\n")
        for i, table in enumerate(content['tables']):
            if len(table) > 0:
                markdown_content.append(f"\n### Table {i+1}\n")
                
                # Create markdown table
                if len(table[0]) > 0:
                    # Header row
                    header = "| " + " | ".join(table[0]) + " |"
                    markdown_content.append(header)
                    
                    # Separator
                    separator = "| " + " | ".join(["---"] * len(table[0])) + " |"
                    markdown_content.append(separator)
                    
                    # Data rows
                    for row in table[1:]:
                        if len(row) == len(table[0]):
                            data_row = "| " + " | ".join(row) + " |"
                            markdown_content.append(data_row)
                
                markdown_content.append("")
    
    return "\n".join(markdown_content)

def analyze_document_complexity(content: Dict) -> Dict:
    """Enhanced document analysis"""
    total_paragraphs = len(content['paragraphs'])
    total_tables = len(content['tables'])
    
    # Calculate complexity metrics
    complexity_scores = [para['complexity_score'] for para in content['paragraphs']]
    avg_complexity = sum(complexity_scores) / len(complexity_scores) if complexity_scores else 0
    
    # Count cleaning opportunities
    all_opportunities = []
    for para in content['paragraphs']:
        all_opportunities.extend(para['cleaning_opportunities'])
    
    opportunity_counts = {}
    for opp in all_opportunities:
        opportunity_counts[opp] = opportunity_counts.get(opp, 0) + 1
    
    # Estimate content density
    total_text_length = sum(para['length'] for para in content['paragraphs'])
    avg_paragraph_length = total_text_length / total_paragraphs if total_paragraphs > 0 else 0
    
    # Count styles
    styles = {}
    for para in content['paragraphs']:
        style = para['style']
        styles[style] = styles.get(style, 0) + 1
    
    complexity = {
        'total_paragraphs': total_paragraphs,
        'total_tables': total_tables,
        'total_text_length': total_text_length,
        'avg_paragraph_length': round(avg_paragraph_length, 2),
        'avg_complexity_score': round(avg_complexity, 2),
        'style_distribution': styles,
        'cleaning_opportunities': opportunity_counts,
        'complexity_score': 'Low' if avg_complexity < 0.3 else 'Medium' if avg_complexity < 0.6 else 'High'
    }
    
    return complexity

def main():
    st.title("🚀 Style Guide Cleaner Pro")
    st.markdown("Upload a DOCX style guide file to clean and optimize it for LLM translation with AI-powered enhancement")
    
    # Sidebar
    st.sidebar.header("🔧 Settings")
    
    # OpenAI API Key input
    st.sidebar.subheader("🤖 OpenAI Integration")
    openai_api_key = st.sidebar.text_input(
        "OpenAI API Key",
        type="password",
        help="Enter your OpenAI API key for AI-powered content optimization",
        value=st.session_state.get('openai_api_key', '')
    )
    
    # Check if API key is available from environment
    env_api_key = os.getenv('OPENAI_API_KEY')
    if env_api_key:
        st.sidebar.info("🔑 API key found in environment")
    
    if openai_api_key:
        st.session_state['openai_api_key'] = openai_api_key
        client = setup_openai()
        openai_available = client is not None
    elif env_api_key:
        # Try using environment variable
        client = setup_openai()
        openai_available = client is not None
    else:
        client = None
        openai_available = False
    
    if openai_available:
        st.sidebar.success("✅ OpenAI connected")
    else:
        st.sidebar.warning("⚠️ OpenAI not connected")
        if not env_api_key and not openai_api_key:
            st.sidebar.info("💡 Add your API key above or create a .env file")
    
    # Processing options
    st.sidebar.subheader("📝 Processing Options")
    output_format = st.sidebar.selectbox(
        "Output Format",
        ["Markdown", "Plain Text"],
        help="Choose the output format for the cleaned content"
    )
    
    # Content removal options
    st.sidebar.subheader("🗑️ Content Removal Options")
    
    remove_tables = st.sidebar.checkbox(
        "Remove Tables",
        value=False,
        help="Remove tables from the output"
    )
    
    remove_urls = st.sidebar.checkbox(
        "Remove URLs & Links",
        value=True,
        help="Remove all web links and URLs"
    )
    
    remove_page_refs = st.sidebar.checkbox(
        "Remove Page References",
        value=True,
        help="Remove page numbers and cross-references"
    )
    
    remove_version_refs = st.sidebar.checkbox(
        "Remove Version References",
        value=True,
        help="Remove version numbers and dates"
    )
    
    remove_metadata = st.sidebar.checkbox(
        "Remove Document Metadata",
        value=True,
        help="Remove document metadata references"
    )
    
    # Human workflow options
    st.sidebar.subheader("👥 Human Workflow Removal")
    
    remove_human_workflow = st.sidebar.checkbox(
        "Remove Human Workflow Sections",
        value=True,
        help="Remove sections about project management, queries, feedback"
    )
    
    remove_contact_instructions = st.sidebar.checkbox(
        "Remove Contact Instructions",
        value=True,
        help="Remove instructions to contact PM, supervisors, etc."
    )
    
    remove_qa_workflows = st.sidebar.checkbox(
        "Remove QA Workflows",
        value=True,
        help="Remove quality assurance processes involving humans"
    )
    
    remove_admin_procedures = st.sidebar.checkbox(
        "Remove Admin Procedures",
        value=True,
        help="Remove system login, database updates, etc."
    )
    
    # Cleaning intensity options
    st.sidebar.subheader("🧹 Cleaning Intensity")
    
    aggressive_cleaning = st.sidebar.checkbox(
        "Aggressive Cleaning",
        value=False,
        help="Apply more aggressive cleaning rules"
    )
    
    remove_wordy_text = st.sidebar.checkbox(
        "Remove Wordy Text",
        value=True,
        help="Remove unnecessary words and phrases"
    )
    
    remove_excessive_punctuation = st.sidebar.checkbox(
        "Remove Excessive Punctuation",
        value=True,
        help="Clean up excessive dots, dashes, etc."
    )
    
    # AI optimization
    st.sidebar.subheader("🤖 AI Optimization")
    
    use_openai_optimization = st.sidebar.checkbox(
        "AI Content Optimization",
        value=openai_available,
        disabled=not openai_available,
        help="Use OpenAI to rewrite content for better conciseness and clarity"
    )
    
    ai_aggressive_mode = st.sidebar.checkbox(
        "AI Aggressive Mode",
        value=False,
        disabled=not openai_available,
        help="AI will be more aggressive about removing human workflow elements"
    )
    
    if use_openai_optimization and not openai_available:
        st.sidebar.error("OpenAI API key required for AI optimization")
    
    # Main content area
    uploaded_file = st.file_uploader(
        "Choose a DOCX file",
        type=['docx'],
        help="Upload a Microsoft Word document (.docx) containing your style guide"
    )
    
    if uploaded_file is not None:
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_file_path = tmp_file.name
            
            # Load document
            doc = Document(tmp_file_path)
            
            # Extract content
            with st.spinner("Processing document..."):
                content = extract_text_from_docx(doc)
                complexity = analyze_document_complexity(content)
            
            st.success(f"✅ Document loaded successfully!")
            
            # Display enhanced analysis
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Paragraphs", complexity['total_paragraphs'])
            with col2:
                st.metric("Tables", complexity['total_tables'])
            with col3:
                st.metric("Complexity", complexity['complexity_score'])
            with col4:
                st.metric("Avg Complexity", f"{complexity['avg_complexity_score']:.2f}")
            
            # Show cleaning opportunities
            if complexity['cleaning_opportunities']:
                st.subheader("🧹 Cleaning Opportunities Found")
                opp_cols = st.columns(len(complexity['cleaning_opportunities']))
                for i, (opp, count) in enumerate(complexity['cleaning_opportunities'].items()):
                    with opp_cols[i]:
                        st.metric(opp.replace('_', ' ').title(), count)
            
            # Show style distribution
            if complexity['style_distribution']:
                st.subheader("📊 Style Distribution")
                style_df = pd.DataFrame(list(complexity['style_distribution'].items()), 
                                      columns=['Style', 'Count'])
                st.bar_chart(style_df.set_index('Style'))
            
            # Create options dictionary from sidebar settings
            cleaning_options = {
                'remove_tables': remove_tables,
                'remove_urls': remove_urls,
                'remove_page_refs': remove_page_refs,
                'remove_version_refs': remove_version_refs,
                'remove_metadata': remove_metadata,
                'remove_human_workflow': remove_human_workflow,
                'remove_contact_instructions': remove_contact_instructions,
                'remove_qa_workflows': remove_qa_workflows,
                'remove_admin_procedures': remove_admin_procedures,
                'remove_wordy_text': remove_wordy_text,
                'remove_excessive_punctuation': remove_excessive_punctuation
            }
            
            # Process content
            if remove_tables:
                content['tables'] = []
            
            if aggressive_cleaning:
                for para in content['paragraphs']:
                    para['text'] = clean_text_for_llm(para['text'], aggressive=True, options=cleaning_options)
            
            # Convert to output format
            if output_format == "Markdown":
                cleaned_content = convert_to_markdown(content, use_openai_optimization, client, cleaning_options, ai_aggressive_mode)
                file_extension = "md"
                mime_type = "text/markdown"
            else:
                # For plain text, apply basic cleaning
                cleaned_paragraphs = []
                for para in content['paragraphs']:
                    cleaned_text = clean_text_for_llm(para['text'], aggressive=aggressive_cleaning, options=cleaning_options)
                    if cleaned_text:
                        cleaned_paragraphs.append(cleaned_text)
                cleaned_content = "\n\n".join(cleaned_paragraphs)
                file_extension = "txt"
                mime_type = "text/plain"
            
            # Display preview
            st.subheader("📖 Cleaned Content Preview")
            with st.expander("View Cleaned Content", expanded=True):
                st.text_area(
                    "Preview",
                    cleaned_content,
                    height=400,
                    disabled=True
                )
            
            # Download button
            st.subheader("💾 Download Cleaned File")
            col1, col2 = st.columns([1, 1])
            with col1:
                st.download_button(
                    label=f"📥 Download {output_format} file",
                    data=cleaned_content,
                    file_name=f"cleaned_style_guide.{file_extension}",
                    mime=mime_type,
                    help="Download the cleaned and optimized style guide file"
                )
            
            with col2:
                if st.button("📊 Download Analysis Report"):
                    analysis_report = json.dumps(complexity, indent=2)
                    st.download_button(
                        label="📥 Download Analysis",
                        data=analysis_report,
                        file_name="document_analysis.json",
                        mime="application/json"
                    )
            
            # Cleanup
            os.unlink(tmp_file_path)
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            st.info("Please make sure you're uploading a valid DOCX file.")
    
    else:
        # Show instructions
        st.info("👆 Please upload a DOCX file to get started")
        
        st.markdown("""
        ## 🚀 Enhanced Features
        
        This enhanced version includes:
        
        ### 🤖 AI-Powered Optimization
        - **Content Rewriting**: Uses OpenAI to make text more concise and clear
        - **Human Workflow Removal**: Automatically removes process elements not useful for LLMs
        - **Smart Analysis**: Identifies complex content that needs optimization
        - **Context-Aware**: Considers document structure and style when optimizing
        
        ### 🧹 Advanced Cleaning
        - **Intelligent Detection**: Identifies specific cleaning opportunities
        - **Human Process Removal**: Removes PM workflows, communication protocols, admin procedures
        - **Complexity Scoring**: Rates content complexity for better optimization
        - **Aggressive Mode**: More thorough cleaning for heavily formatted documents
        
        ### 📊 Enhanced Analytics
        - **Detailed Analysis**: Comprehensive document complexity metrics
        - **Cleaning Opportunities**: Shows what will be cleaned (including human workflows)
        - **Export Reports**: Download detailed analysis as JSON
        
        ### 🎯 What Gets Removed for LLMs:
        - ❌ Process instructions ("consult with project manager")
        - ❌ Human workflows ("send to reviewer", "get approval from client")
        - ❌ Communication protocols ("email the team", "call the client")
        - ❌ Administrative procedures ("log in the system", "update database")
        - ❌ Quality assurance workflows involving humans
        - ❌ References to human roles for process purposes
        
        ### ✅ What Gets Optimized:
        - ✅ Translation rules and guidelines
        - ✅ Terminology preferences
        - ✅ Style conventions and grammar rules
        - ✅ Cultural considerations
        - ✅ Brand voice guidelines
        - ✅ Technical specifications for translation
        
        ## 🔧 Setup Instructions
        
        1. **Get OpenAI API Key**: Visit [OpenAI Platform](https://platform.openai.com/api-keys)
        2. **Enter API Key**: Add your key in the sidebar
        3. **Upload Document**: Choose your DOCX file
        4. **Configure Options**: Adjust cleaning and optimization settings
        5. **Download Result**: Get your optimized style guide
        
        ## 💡 Tips for Best Results
        
        - Use OpenAI optimization for complex or wordy documents
        - Enable aggressive cleaning for heavily formatted files
        - Review the analysis report to understand what was cleaned
        - The AI optimization works best on paragraphs with complexity score > 0.3
        """)

if __name__ == "__main__":
    import pandas as pd
    main()
